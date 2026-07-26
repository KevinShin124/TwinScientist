"""Generate the comprehensive TwinScientist summary document as .docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('TwinScientist')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('基于国产开源大模型的自主科研智能体')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

doc.add_paragraph()

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run('技术架构 · 核心亮点 · 竞品对比 · 优化路线')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'版本: 2.0 (Optimized)').font.size = Pt(11)
meta.add_run(f'\n日期: {datetime.now().strftime("%Y年%m月%d日")}').font.size = Pt(11)
meta.add_run('\n赛事: 挑战杯 XH-202619 "基于国产开源大模型的AI Scientist的研发与应用"').font.size = Pt(11)
meta.add_run('\n模型: Qwen 系列 (阿里云百炼平台)').font.size = Pt(11)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading('目录', level=1)
toc_items = [
    '一、项目概述与定位',
    '二、系统架构总览',
    '三、核心认知流水线 (14个节点)',
    '四、六大升级优化 (v2.0)',
    '五、国际竞品对比分析',
    '六、核心差异化优势',
    '七、技术指标与评分对标',
    '八、比赛提交清单',
    '九、未来演进方向',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ============================================================
# SECTION 1: Overview
# ============================================================
doc.add_heading('一、项目概述与定位', level=1)

doc.add_paragraph(
    'TwinScientist 是一款面向环境—人体关联研究的自主科研与实验迭代智能体。'
    '系统融合温湿度、CO₂、PPG、血氧、HRV 及视觉疲劳等多源真实传感器数据，'
    '自动完成数据清洗、时间对齐、质量评估、科学假设生成与实验方案设计，'
    '并根据实测结果持续修正假设和优化下一轮实验。'
)

doc.add_heading('核心定位', level=2)
bullets = [
    '领域聚焦: 环境—人体关联 (Environment-Human Association)',
    '数据驱动: 16个真实传感器CSV + 8个生物特征CSV，非合成数据',
    '方法论: 因果推断 (CCM/Granger/Counterfactual)，非简单相关性',
    '个体化: N-of-1 研究范式，同一受试者跨场景对比分析',
    '自主闭环: 从问题输入到标准化报告输出的全自动科研流程',
    '人机协同: 伦理审查 + 五维评审 + 辩论 + 人在回路审核',
    '国产化: 基于 Qwen 系列大模型 + 阿里云百炼平台，完全自主可控',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# ============================================================
# SECTION 2: Architecture
# ============================================================
doc.add_heading('二、系统架构总览', level=1)

doc.add_paragraph(
    'TwinScientist 采用 LangGraph 认知图编排架构，14个专用认知节点组成完整的科研流水线。'
    '每个节点对应一个科学方法论步骤，通过 AgentState 统一状态空间进行数据流转。'
    '架构对标 Google DeepMind Co-Scientist 的 6-Agent 设计范式，'
    '同时针对比赛要求增加了辩论、教学注释和人在回路机制。'
)

doc.add_heading('技术栈', level=2)
tech_table = doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
tech_data = [
    ('编排框架', 'LangGraph (StateGraph + MemorySaver)'),
    ('大语言模型', 'Qwen-Max (阿里云百炼平台, OpenAI兼容API)'),
    ('因果推断', 'CCM收敛交叉映射 / Granger因果检验 / 反事实推演'),
    ('文献检索', 'Crossref REST API + arXiv API + Semantic Scholar'),
    ('知识图谱', 'NetworkX 有向图 + 实体-关系自动提取'),
    ('假设推理', 'LogicEngine: 归纳/演绎/溯因三路推理'),
    ('数据管道', 'Daltons 格式解析 + 时间序列对齐 + 质量评估'),
    ('输出格式', 'Markdown 标准化报告 + 可视化图表注入'),
]
for i, (k, v) in enumerate(tech_data):
    tech_table.cell(i, 0).text = k
    tech_table.cell(i, 1).text = v

doc.add_heading('认知流水线拓扑', level=2)
doc.add_paragraph(
    'START → 伦理审查 → 文献调研 → 假设生成 → 淘汰赛 → 实验设计 → '
    '数据分析 → 结果解读 → 五维评审 → [反思修正 ↻] → 智能体辩论 → '
    '终止评估 → 报告撰写 → PI总结 → 人机审核 → 自我进化 → END'
)

doc.add_page_break()

# ============================================================
# SECTION 3: 14 Nodes
# ============================================================
doc.add_heading('三、核心认知流水线 (14 个节点)', level=1)

nodes = [
    ('1. 伦理审查 (Ethics Check)',
     '第一道安全防线。结构化 JSON 输出解析，三段式评估 (BLOCKED/HUMAN_REVIEW/APPROVED)。'
     '对标 Anthropic Deep Research 的安全评估模式。'),
    ('2. 文献调研 (Literature Review)',
     '并行搜索 Crossref + arXiv，LLM 提取 >=8 条结构化科学事实，'
     'CitationValidator 交叉验证 DOI/PMID，自动构建知识图谱。'),
    ('3. 假设生成 (Hypothesis Generation)',
     'LogicEngine 三路推理 (归纳/演绎/溯因) + LLM 上下文增强，'
     '合并去重 + 逻辑一致性检查，每轮生成 10-16 个候选假设。'),
    ('4. 淘汰赛评估 (Tournament Eval)',
     'N 选 1 淘汰赛制，两两对比，LLM 裁判打分。'
     '胜出假设进入实验设计，淘汰者标记 pruned。'),
    ('5. 实验设计 (Experiment Design)',
     '基于优胜假设设计完整实验方案，包含 12 个标准字段。'
     '自动关联传感器数据文件路径。'),
    ('6. 数据分析 (Data Analysis)',
     '自动选择最优因果推断方法 (Granger/CCM/Counterfactual)，'
     '加载真实传感器数据，构建标准化 EvidenceEntry。'),
    ('7. 结果解读 (Interpretation)',
     'LLM 综合分析结果，更新 Bayesian 后验概率，'
     '计算收敛度 (1 - max_confidence_change)。'),
    ('8. 五维评审 (Reviewer Agent)',
     '新颖性/可行性/方法论/证据/影响力 五维打分，>=60 通过，<60 打回。'
     'Bayesian log-odds 置信度更新。'),
    ('9. 反思修正 (Reflection)',
     '根因分析 → 派生修正性假设，失败资产化存储。'
     '自动剪枝低置信度/无子节点的假设。'),
    ('10. 智能体辩论 (Debate)',
     'Pro/Con/Judge 三方对抗辩论，1 轮精简模式。'
     '辩论结果写入假设置信度，记录完整辩论历史。'),
    ('11. 终止评估 (Termination Eval)',
     '四维收敛评估: 语义相似度 + 方法论稳定性 + 证据覆盖度 + 假设空间聚焦。'
     '跨学科迁移分析。'),
    ('12. 报告撰写 (Report Writing)',
     '生成 13 节标准化《科学假设与研究计划》，'
     '程序化组装真实数据 + LLM 生成 Rationale/Abstract，'
     '自动注入可视化图表。'),
    ('13. PI 总结 (PI Agent Meeting)',
     '整合多智能体成果，保护真实因果推断数据不被覆盖。'),
    ('14. 人机审核 + 自我进化 (HITL + Evolution)',
     'CLI 自动确认 / UI 人工审核双模式。'
     '提取成功模式共性和失败规律，蒸馏 meta-insights。'),
]
for title, desc in nodes:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

doc.add_page_break()

# ============================================================
# SECTION 4: 6 Upgrades
# ============================================================
doc.add_heading('四、六大升级优化 (v2.0)', level=1)

doc.add_paragraph(
    '以下升级基于对 OpenAI Deep Research、Google DeepMind Co-Scientist、'
    'Sakana AI Scientist v2、Anthropic Computer Use 等国际前沿 AI Scientist 系统的'
    '深度分析，结合 Apple/Google 级工程标准实施。'
)

upgrades = [
    ('Upgrade 1: Streaming + Live Dashboard (对标 OpenAI Deep Research)',
     '核心模块', 'core/progress.py',
     '用 LangGraph astream_events() 替代 ainvoke()，实现每个认知节点的实时进度展示。'
     '包含状态指示器、耗时统计、关键指标提取 (文献数/假设数/因果强度/评审分数)。'
     '评委在演示视频中可以看到流水线逐节点推进，而非干等 5-10 分钟后一次性看到结果。'
     '\n\n技术亮点: 零开销事件捕获 (LangGraph 原生 streaming)，ASCII-safe 兼容 Windows 控制台。'),
    ('Upgrade 2: Multi-Modal Output (对标 Apple 级产品体验)',
     '核心模块', 'core/visualization.py',
     '自动生成 5 种可视化图表并注入报告:'
     '\n  - 因果推断 ASCII 图表 (方法/强度/方向/统计依据)'
     '\n  - 假设树 Mermaid 思维导图'
     '\n  - 证据时间线 (实验→方法→证据强度→结论)'
     '\n  - 收敛曲线 (迭代轮次 vs 收敛度)'
     '\n  - 数据质量概览表'
     '\n\n技术亮点: 纯 Python 实现，零外部依赖，ASCII/Mermaid 格式兼容所有 Markdown 渲染器。'),
    ('Upgrade 3: Adaptive Iterations (对标 OpenAI Test-Time Compute Scaling)',
     '核心模块', 'core/adaptive.py',
     '6 个复杂度信号 (交互项/因果/时序/多变量/个性化/跨领域) 加权计算，'
     '动态分配迭代预算: 简单问题 3 轮，中等 5 轮，复杂机制问题 15+ 轮。'
     '不是死板的固定迭代，而是根据问题难度智能调度计算资源。'
     '\n\n技术亮点: 加权复杂度评分算法，中英文双语关键词匹配，领域加成。'),
    ('Upgrade 4: Specialized Agent Personas (对标 Google Co-Scientist 6-Agent)',
     '核心模块', 'core/prompts.py',
     '5 个专用 Agent 角色，每个有独立的定位、评估标准和输出格式:'
     '\n  - Generation Agent: 假设生成专家 (创意引擎)'
     '\n  - Reflection Agent: 反思审查专家 (魔鬼代言人)'
     '\n  - Ranking Agent: 淘汰赛裁判 (竞技场裁判)'
     '\n  - Proximity Agent: 文献新颖性检查 (原创性守护者)'
     '\n  - Meta-Review Agent: 综合评审专家 (总编辑)'
     '\n\n技术亮点: 与 Google 的 6-Agent 架构一一对应，但针对 Qwen 的中文指令遵循偏好深度定制。'),
    ('Upgrade 5: Research Memory (对标 Google 跨会话知识持久化)',
     '核心模块', 'core/memory.py',
     '每次研究完成后自动提取关键发现 (假设/证据/结论) 存入 JSON 知识库。'
     '新问题自动检索相关历史发现并注入 LLM prompt，实现"越用越聪明"。'
     '支持关键词匹配检索、证据质量加权、领域过滤。'
     '\n\n技术亮点: 零外部依赖 (纯 JSON 文件)，可在生产环境无缝升级到向量数据库。'),
    ('Upgrade 6: SFT Data Pipeline (比赛允许的微调数据收集)',
     '核心模块', 'core/sft_pipeline.py',
     '自动从每次研究会话中收集 5 类 (prompt, completion) 训练对:'
     '\n  - 假设生成: 文献事实 → 结构化假设'
     '\n  - 方法选择: 数据特征 → 因果推断方法'
     '\n  - 同行评审: 假设 → 五维评分'
     '\n  - Rationale 撰写: 研究上下文 → 解决思路'
     '\n  - Abstract 撰写: 研究结果 → 论文摘要'
     '\n导出标准 JSONL 格式，可直接用于 Qwen SFT 微调。'
     '\n\n技术亮点: 比赛规则明确允许 SFT/微调，此管道实现训练数据的自动化收集。'),
]
for title, module_label, module_path, desc in upgrades:
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    p.add_run(f'{module_label}: ').bold = True
    p.add_run(module_path)
    doc.add_paragraph(desc)

doc.add_page_break()

# ============================================================
# SECTION 5: Competitive Analysis
# ============================================================
doc.add_heading('五、国际竞品对比分析', level=1)

comp_table = doc.add_table(rows=7, cols=8, style='Light Grid Accent 1')
comp_headers = ['系统', '领域', '数据源', '方法论', '因果推断', '辩论', '人在回路', 'N-of-1']
for i, h in enumerate(comp_headers):
    comp_table.cell(0, i).text = h
    for p in comp_table.cell(0, i).paragraphs:
        for run in p.runs:
            run.bold = True

comp_data = [
    ['Sakana AI\nScientist v2', 'ML/AI\n通用', '代码执行\n结果', '模板+LLM', '❌ 无', '❌ 自动\n打分', '❌ 无', '❌'],
    ['Agent\nLaboratory', 'ML/NLP\n通用', '代码执行\n结果', '多Agent\n协作', '❌ 无', '❌', '可选\n但弱', '❌'],
    ['GPT\nResearcher', '通用Web\n调研', '网页搜索', 'Plan+\nSolve', '❌ 无', '❌', '❌', '❌'],
    ['OpenAI Deep\nResearch', '通用\n研究', '浏览器+\nPython', 'RL微调\no3', '❌ 无', '❌', '❌', '❌'],
    ['Google\nCo-Scientist', '生物医学\n药物发现', '文献库+\n计算', '6-Agent\n协作', '❌ 无', '❌', '可选', '❌'],
    ['★ Twin-\nScientist ★', '环境-人体\n关联', '真实IoT\n传感器', '因果推断\n+LLM', '✅ CCM\nGranger\nCounterfactual', '✅ Pro/Con\n/Judge', '✅ 伦理+\n评审+\n对话', '✅ 跨场景\n个体化\n对比'],
]
for i, row_data in enumerate(comp_data):
    for j, val in enumerate(row_data):
        comp_table.cell(i+1, j).text = val

doc.add_paragraph()
doc.add_paragraph(
    '关键发现: 所有国际竞品均局限于"代码执行结果"或"网页搜索"，无法接入真实物理传感器数据。'
    '没有竞品同时具备因果推断 + 辩论 + 人在回路 + N-of-1 能力。'
    'TwinScientist 在"环境-人体关联"这一垂直领域具有不可替代的差异化优势。'
).italic = True

doc.add_page_break()

# ============================================================
# SECTION 6: Differentiators
# ============================================================
doc.add_heading('六、核心差异化优势', level=1)

advantages = [
    ('🔥 真实物理传感器数据 (最大壁垒)',
     '16 个环境传感器 CSV + 8 个生物特征 CSV，来自真实家庭环境 (H1 卧室/厨房/客厅/书房)。'
     '所有竞品只能跑代码实验或搜索网页。AI-Scientist 永远做不到"加载 H1_Bedroom.csv, '
     '18000 条 Daltons 记录, Granger 检验 CO₂→HRV 因果方向"。'),
    ('🔥 因果推断方法论 (非相关性)',
     'CCM 收敛交叉映射 + Granger 因果检验 + 反事实推演。'
     '不是"数据看起来相关"，而是"因果方向 X→Y, 证据强度 0.89, p<0.01"。'
     '每条证据附带完整统计依据 (p值/效应量/置信区间)。'),
    ('🔥 N-of-1 个体化研究 (前沿方向)',
     'N-of-1 + LLM + IoT 传感器是 2025 年 Nature Digital Medicine 和 '
     'The Lancet Digital Health 关注的前沿方向。目前无成熟开源框架。'
     'TwinScientist 天然支持同一受试者跨场景对比分析。'),
    ('🔥 可追溯证据链',
     'EvidenceEntry 标准化结构: 方法参数 (CCM library size, Granger lag)、'
     '统计依据 (p-value, F-statistic, rho)、验证结果 (收敛性检查, 平稳性检查)。'
     '每条结论都可追溯到原始数据和具体方法。'),
    ('🔥 智能体辩论 + 人在回路 (比赛硬性要求)',
     'Pro/Con/Judge 三方对抗辩论 + 伦理审查 + 五维评审 + 人机审核。'
     '教学注释自动生成，每个节点解释科学原理。'
     '竞品中无一同时具备所有这些能力。'),
    ('🔥 国产化技术栈 (比赛硬性要求)',
     'Qwen 系列大模型 + 阿里云百炼平台 + LangGraph 编排。'
     '完全自主可控，符合比赛"基于国产开源大模型"的核心要求。'),
]
for title, desc in advantages:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

doc.add_page_break()

# ============================================================
# SECTION 7: Scoring
# ============================================================
doc.add_heading('七、技术指标与评分对标', level=1)

doc.add_paragraph(
    '以下为比赛评分标准 (总分 100) 与 TwinScientist 的对应能力映射:'
)

score_table = doc.add_table(rows=5, cols=4, style='Light Grid Accent 1')
score_headers = ['评分维度', '分值', '评分细则', 'TwinScientist 对应能力']
for i, h in enumerate(score_headers):
    score_table.cell(0, i).text = h
    for p in score_table.cell(0, i).paragraphs:
        for run in p.runs:
            run.bold = True

score_data = [
    ['科学价值', '40', '创新性 (0-20)\n逻辑自洽性 (0-20)', '因果推断方法论创新\nLogicEngine 三路推理 + LLM 增强\nBayesian 置信度更新\n假设树动态生长与剪枝'],
    ['技术实现', '30', 'Agent 协作设计 (0-15)\n多模态处理 (0-15)', '14 节点 LangGraph 编排\n6 大升级 (Streaming/可视化/自适应/Persona/记忆/SFT)\nDaltons 格式多源数据融合\n知识图谱自动构建'],
    ['应用潜力', '30', '实际场景支撑 (0-10)\n转化潜力 (0-10)\n社会价值 (0-10)', 'N-of-1 个性化环境健康\n真实传感器数据验证\n可追溯证据链\nSFT 数据管道支撑持续优化'],
    ['总分', '100', '', '全面覆盖所有评分维度\n垂直领域壁垒 (真实数据)\n国际对标架构设计'],
]
for i, row_data in enumerate(score_data):
    for j, val in enumerate(row_data):
        score_table.cell(i+1, j).text = val

doc.add_paragraph()

doc.add_heading('关键性能指标', level=2)
metrics = [
    '文献验证率: 8% → 优化后预期 >50% (修复 arXiv 301, Crossref 400)',
    '单次研究耗时: 5-10 分钟 (含 3-15 轮自适应迭代)',
    '假设生成: 每轮 10-16 个候选假设 (LogicEngine + LLM)',
    '因果推断方法: 3 种 (Granger/CCM/Counterfactual) 自动选择',
    '辩论 LLM 调用: 从 9 次降至 3 次 (1 轮精简)',
    '教学注释: 6 个关键节点自动生成中文解释',
    '代码规模: 核心逻辑 ~2,500 行 + 5 个新模块 ~1,100 行',
]
for m in metrics:
    doc.add_paragraph(m, style='List Bullet')

doc.add_page_break()

# ============================================================
# SECTION 8: Submission Checklist
# ============================================================
doc.add_heading('八、比赛提交清单', level=1)

doc.add_paragraph('根据比赛要求，需提交以下材料 (截止 2026年9月5日):')

checklist = [
    ('PDF 文档 (≤20页)', '本文件即为最终文档的基础。包含: 研究背景、AI Scientist 架构设计、'
     '核心功能讲解、实验验证案例、N-of-1 分析结果。'),
    ('演示视频 (≤10分钟)', '建议叙事: 问题引入 (2min) → 系统演示 (3min) → 结果展示 (3min) → 架构讲解 (2min)。'
     'Streaming Dashboard 提供实时进度展示。'),
    ('源代码', '完整 Python 项目，含 LangGraph 编排、14 个认知节点、6 大升级模块。'
     '依赖: langgraph, httpx, networkx, numpy, pandas, pydantic-settings。'),
    ('数据文件', '16 个传感器 CSV + 8 个生物特征 CSV。脱敏后提交。'),
    ('项目简介页面', '比赛官网要求的前端展示页面。'),
    ('10 分钟演示视频', '建议使用 Gradio UI 录制，展示完整流水线运行。'),
]
for title, desc in checklist:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

doc.add_page_break()

# ============================================================
# SECTION 9: Future
# ============================================================
doc.add_heading('九、未来演进方向', level=1)

doc.add_paragraph(
    '以下为比赛后的长期演进路线，对标国际一流 AI Scientist 系统的发展方向:'
)

future = [
    ('SFT 微调 (短期)',
     '利用 SFT Pipeline 收集的数据对 Qwen 进行领域微调，'
     '提升假设生成、因果方法选择和同行评审的准确性。'),
    ('向量数据库升级 (中期)',
     '将 Research Memory 从 JSON 文件升级为向量数据库 (Milvus/Chroma)，'
     '实现语义相似度检索，支持大规模知识积累。'),
    ('多模态数据融合 (中期)',
     '接入视觉疲劳数据 (眼动追踪 + 面部表情)，'
     '实现环境 × 生理 × 视觉三模态因果推断。'),
    ('RL 策略优化 (长期)',
     '用 MC 经验库的累积数据训练路由策略网络，'
     '替代当前的确定性路由，实现真正的 LLM 动态编排。'),
    ('多受试者扩展 (长期)',
     '从 N-of-1 扩展到 N-of-N，支持群体水平的环境—健康因果推断，'
     '同时保持个体化分析能力。'),
    ('开源社区 (长期)',
     '将 N-of-1 数据集打包为标准化 benchmark，'
     '推动 AI Scientist 领域在真实传感器数据上的基准测试。'),
]
for title, desc in future:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

# ============================================================
# FOOTER
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— TwinScientist v2.0 —')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = footer2.add_run('基于国产开源大模型 (Qwen) 的自主科研智能体')
run2.font.size = Pt(9)
run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
output_path = r'C:\Users\KevinShin\Desktop\TwinScientist_技术文档_v2.0.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')